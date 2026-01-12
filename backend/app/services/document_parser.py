"""
Document parsing service for extracting text from PDF and DOCX files.
"""

import logging
from pathlib import Path
from typing import Optional
import pypdf
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParserError(Exception):
    """Custom exception for document parsing errors."""
    pass


class DocumentParser:
    """Service for parsing PDF and DOCX documents."""

    @staticmethod
    async def extract_text(file_path: str, file_type: str) -> str:
        """
        Extract text from a document file.

        Args:
            file_path: Path to the document file
            file_type: File extension (pdf or docx)

        Returns:
            Extracted text content

        Raises:
            DocumentParserError: If extraction fails
        """
        try:
            if file_type.lower() == "pdf":
                return await DocumentParser._extract_from_pdf(file_path)
            elif file_type.lower() == "docx":
                return await DocumentParser._extract_from_docx(file_path)
            else:
                raise DocumentParserError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise DocumentParserError(f"Failed to extract text: {str(e)}")

    @staticmethod
    async def _extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file using pypdf with pdfplumber fallback.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text

        Raises:
            DocumentParserError: If extraction fails
        """
        text_parts = []

        try:
            # First try with pypdf (faster)
            logger.info(f"Attempting to extract text from PDF using pypdf: {file_path}")
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")

            # If pypdf extracted reasonable amount of text, return it
            full_text = "\n\n".join(text_parts)
            if len(full_text.strip()) > 100:  # At least 100 characters
                logger.info(f"Successfully extracted {len(full_text)} characters from PDF using pypdf")
                return full_text

            # If pypdf didn't extract much, try pdfplumber (handles complex layouts better)
            logger.info("pypdf extracted minimal text, trying pdfplumber...")
            text_parts = []

        except Exception as e:
            logger.warning(f"pypdf extraction failed: {str(e)}, falling back to pdfplumber")

        # Fallback to pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise DocumentParserError("No text could be extracted from PDF. The file may be scanned or image-based.")

            logger.info(f"Successfully extracted {len(full_text)} characters from PDF using pdfplumber")
            return full_text

        except Exception as e:
            logger.error(f"pdfplumber extraction also failed: {str(e)}")
            raise DocumentParserError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    async def _extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text

        Raises:
            DocumentParserError: If extraction fails
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            # Also extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_text.append(row_text)

            # Combine paragraphs and tables
            full_text = "\n\n".join(paragraphs)
            if table_text:
                full_text += "\n\n--- Tables ---\n" + "\n".join(table_text)

            if not full_text.strip():
                raise DocumentParserError("No text could be extracted from DOCX file")

            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise DocumentParserError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def validate_file(file_path: str, max_size_bytes: int = 50 * 1024 * 1024) -> bool:
        """
        Validate that a file exists and is within size limits.

        Args:
            file_path: Path to file
            max_size_bytes: Maximum allowed file size (default 50MB)

        Returns:
            True if file is valid

        Raises:
            DocumentParserError: If file is invalid
        """
        path = Path(file_path)

        if not path.exists():
            raise DocumentParserError(f"File not found: {file_path}")

        if not path.is_file():
            raise DocumentParserError(f"Path is not a file: {file_path}")

        file_size = path.stat().st_size
        if file_size > max_size_bytes:
            size_mb = file_size / (1024 * 1024)
            max_mb = max_size_bytes / (1024 * 1024)
            raise DocumentParserError(
                f"File too large: {size_mb:.1f}MB (max: {max_mb:.1f}MB)"
            )

        return True
